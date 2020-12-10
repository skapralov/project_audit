import io
import re
import time

import requests
from lxml import html
import pygsheets
from docx import Document
from docx.shared import Cm
from google.oauth2 import service_account
from googleapiclient.http import MediaIoBaseDownload
from googleapiclient.discovery import build


SPREADSHEET_URL = 'https://docs.google.com/spreadsheets/d/11irhn_QwAxgy0MZhU5q3OkIWgGk0QffnMnli2qf5OEc/edit?usp=sharing'
SCOPES = ['https://www.googleapis.com/auth/drive']
SERVICE_ACCOUNT_FILE = 'project-audit-296419-6a7ea20a5c8b.json'


def read_google_sheet():
    client = pygsheets.authorize(service_file=SERVICE_ACCOUNT_FILE)
    sheet = client.open_by_url(SPREADSHEET_URL)
    wks = sheet.sheet1
    return wks


def get_data_from_sheet(wks, project_id):
    for row in wks:
        if row[0] == project_id:
            return row
    print('project not found')


def complete_document(data, project_id):
    data = {f'<<{i[0]}>>': i[1] for i in enumerate(data)}
    document = Document('./template.docx')
    document = set_title_project(document, project_id)
    for paragraph in document.paragraphs:
        for key, value in data.items():
            if not str(value).strip():
                continue

            if key in paragraph.text and '<<image>>' in paragraph.text:
                paragraph.text = None
                image = get_image_from_lightshot(value)
                paragraph.runs[-1].add_break()
                paragraph.runs[-1].add_picture(image, width=Cm(16.5))
            elif key in paragraph.text:
                paragraph.text = value

    clean_document = clear_from_tags(document)
    file_name = f'{project_id}_Аудит.docx'
    clean_document.save(file_name)
    print(f'saved file {file_name}')


def set_title_project(document, project_id):
    for paragraph in document.paragraphs:
        for run in paragraph.runs:
            if '<<project>>' in run.text:
                run.text = project_id
                return document


def get_image_from_drive(link):
    print('get image from drive -- ', end='')
    credentials = service_account.Credentials.from_service_account_file(SERVICE_ACCOUNT_FILE, scopes=SCOPES)
    service = build('drive', 'v3', credentials=credentials)
    file_id = link.split('/d/')[1].split('/')[0]
    request = service.files().get_media(fileId=file_id)
    fh = io.BytesIO()
    downloader = MediaIoBaseDownload(fh, request)
    while True:
        status, done = downloader.next_chunk()
        if done:
            print('successfully')
            return fh


def get_image_from_lightshot(link):
    print('get image from lightshot -- ', end='')
    user_agents = {'User-agent': 'Mozilla/5.0 (X11; Linux x86_64; rv:70.0) Gecko/20100101 Firefox/70.0'}
    response = requests.get(link, headers=user_agents)
    document = html.fromstring(response.content)
    src = document.xpath('//meta[@property="og:image"]')[0].get('content')
    response = requests.get(src)
    image = io.BytesIO(response.content)
    print('successfully')
    return image


def clear_from_tags(document):
    for paragraph in document.paragraphs:
        if re.search(r'<<\w{1,5}>>', paragraph.text):
            p = paragraph._element
            p.getparent().remove(p)
            p._p = p._element = None
    return document


def main():
    project_id = input('enter project id: ')
    print('read google sheet -- ', end='')
    wks = read_google_sheet()
    print('successfully')
    data = get_data_from_sheet(wks, project_id)
    complete_document(data, project_id)


if __name__ == '__main__':
    try:
        main()
    except Exception as error:
        print(error)
        time.sleep(999)
